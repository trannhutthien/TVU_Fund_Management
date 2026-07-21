import sys, os, zipfile
sys.stdout.reconfigure(encoding='utf-8')
from lxml import etree

nsmap = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
W = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'

base = r'D:\TotNghiep\TVUDevelopmentFundManager\TVU_Fund_Management\docs\docs\unpacked_word'
out = r'D:\TotNghiep\TVUDevelopmentFundManager\TVU_Fund_Management\docs\docs\110122162_TranNhutThien_KLTN.docx'

# Unpack
doc_path = os.path.join(base, 'word', 'document.xml')
tree = etree.parse(doc_path)
root = tree.getroot()

# Find the body element
body = root.find(f'{W}body')

# Find all <w:p> elements in body
all_paragraphs = body.findall(f'{W}p')
print(f'Total <w:p> in body: {len(all_paragraphs)}')

# Find problematic: <w:p> that has <w:pPr> containing child <w:p> or <w:tbl>
fixes = 0
for p in all_paragraphs:
    ppr = p.find(f'{W}pPr')
    if ppr is None:
        continue
    
    # Check for nested <w:p> or <w:tbl> inside <w:pPr>
    nested = []
    for child in ppr:
        tag = etree.QName(child).localname
        if tag in ('p', 'tbl'):
            nested.append(child)
    
    if not nested:
        continue
    
    # This <w:p> has invalid nesting
    # Strategy: the nested elements are the REAL content
    # Remove the outer <w:p> and insert the nested elements as siblings before it
    parent = p.getparent()
    idx = list(parent).index(p)
    
    # Get the pPr style from each nested <w:p> for reference
    for n in nested:
        # Move each nested element to be a sibling of the outer <w:p>
        parent.insert(idx, n)
        idx += 1
        fixes += 1
    
    # Remove the outer <w:p> (which is now empty of content)
    parent.remove(p)

print(f'Fixed {fixes} nested elements')

# Verify
remaining = 0
for p in body.findall(f'{W}p'):
    ppr = p.find(f'{W}pPr')
    if ppr is None:
        continue
    for child in ppr:
        tag = etree.QName(child).localname
        if tag in ('p', 'tbl'):
            remaining += 1
            print(f'  STILL NESTED: {tag}')

print(f'Remaining issues: {remaining}')

# Write back
tree.write(doc_path, xml_declaration=True, encoding='UTF-8', standalone=True)

# Verify with python-docx
from docx import Document
try:
    doc = Document(out)
    print(f'python-docx: OK, {len(doc.paragraphs)} paragraphs')
    for idx, p in enumerate(doc.paragraphs[:5]):
        txt = p.text[:100] if p.text else '(empty)'
        print(f'  [{idx}] {txt}')
except Exception as e:
    print(f'python-docx ERROR: {e}')

print('Done!')
