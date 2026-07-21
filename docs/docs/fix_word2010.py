import sys, re, os, zipfile
sys.stdout.reconfigure(encoding='utf-8')

base = r'D:\TotNghiep\TVUDevelopmentFundManager\TVU_Fund_Management\docs\docs\unpacked_word'
out = r'D:\TotNghiep\TVUDevelopmentFundManager\TVU_Fund_Management\docs\docs\110122162_TranNhutThien_KLTN.docx'

doc_path = os.path.join(base, 'word', 'document.xml')
with open(doc_path, 'r', encoding='utf-8') as f:
    lines = f.readlines()

print(f'Original lines: {len(lines)}')

# Strategy: Process line by line, track <w:p> and <w:pPr> nesting
# When we find <w:p> inside <w:pPr>, extract it and move it outside

output_lines = []
i = 0
fixes = 0

while i < len(lines):
    line = lines[i]
    stripped = line.strip()
    
    # Check if this line is <w:pPr> that contains nested <w:p> or <w:tbl>
    if stripped == '<w:pPr>':
        # Look ahead to see if there are nested <w:p> or <w:tbl> inside
        j = i + 1
        has_nested = False
        depth = 1
        while j < len(lines) and depth > 0:
            s = lines[j].strip()
            if '<w:pPr>' in s:
                depth += 1
            elif '</w:pPr>' in s:
                depth -= 1
            if depth == 1 and ('<w:p>' in s or '<w:p ' in s or '<w:tbl>' in s):
                has_nested = True
            j += 1
        
        if has_nested:
            # This pPr has nested elements - need to fix
            # Collect the entire outer <w:p> block
            # Find the opening <w:p> that contains this <w:pPr>
            # The outer <w:p> should be on the line before i
            
            # Check if the previous line is the outer <w:p>
            outer_p_line = i - 1
            while outer_p_line >= 0 and lines[outer_p_line].strip() == '':
                outer_p_line -= 1
            
            if outer_p_line >= 0 and ('<w:p>' in lines[outer_p_line] or '<w:p ' in lines[outer_p_line]):
                # Found the outer <w:p> - skip it entirely
                # Skip until we find the matching </w:p>
                skip_start = outer_p_line
                skip_depth = 1
                k = i + 1
                while k < len(lines) and skip_depth > 0:
                    s = lines[k].strip()
                    if '<w:p>' in s or '<w:p ' in s:
                        skip_depth += 1
                    if '</w:p>' in s:
                        skip_depth -= 1
                    k += 1
                
                # Now extract the content from inside <w:pPr>
                # We need to get all the nested <w:p> and <w:tbl> elements
                inner_content = []
                depth = 1
                j = i + 1
                while j < len(lines) and depth > 0:
                    s = lines[j].strip()
                    if '<w:pPr>' in s and '</w:pPr>' not in s:
                        depth += 1
                    elif '</w:pPr>' in s:
                        depth -= 1
                    
                    if depth == 1:
                        # This line is a direct child of pPr
                        if '<w:p>' in s or '<w:p ' in s or '<w:tbl>' in s:
                            # Collect this entire element
                            elem_depth = 1
                            elem_lines = [lines[j]]
                            m = j + 1
                            while m < len(lines) and elem_depth > 0:
                                ms = lines[m].strip()
                                if '<w:p>' in ms or '<w:p ' in ms or '<w:tbl>' in ms:
                                    elem_depth += 1
                                if '</w:p>' in ms or '</w:tbl>' in ms:
                                    elem_depth -= 1
                                elem_lines.append(lines[m])
                                m += 1
                            inner_content.append(elem_lines)
                            j = m - 1  # will be incremented at end of loop
                    j += 1
                
                # Skip the outer <w:p> and its <w:pPr>
                # Don't output the outer <w:p> line
                # Don't output the <w:pPr> line
                # Don't output </w:pPr> line
                # Don't output </w:p> line (outer closing)
                
                # Skip lines from outer_p_line to end of outer block
                i = skip_start
                
                # Output all extracted inner elements
                for elem_lines in inner_content:
                    for el in elem_lines:
                        output_lines.append(el)
                    fixes += 1
                
                # Skip past the outer block
                i = skip_start
                # Find the end of the outer <w:p>
                skip_depth = 0
                while i < len(lines):
                    s = lines[i].strip()
                    if i == skip_start:
                        skip_depth = 1
                        if '<w:p>' in s and '</w:p>' not in s:
                            skip_depth = 1
                        elif '<w:p ' in s and '</w:p>' not in s:
                            skip_depth = 1
                    else:
                        if '<w:p>' in s or '<w:p ' in s:
                            skip_depth += 1
                        if '</w:p>' in s:
                            skip_depth -= 1
                    i += 1
                    if skip_depth <= 0:
                        break
                continue
    
    output_lines.append(line)
    i += 1

print(f'Fixed {fixes} nested elements')
print(f'New lines: {len(output_lines)}')

# Write the fixed file
with open(doc_path, 'w', encoding='utf-8') as f:
    f.writelines(output_lines)

# Verify
content2 = ''.join(output_lines)
in_ppr = False
issues = 0
for i2, l in enumerate(output2_lines := content2.split('\n')):
    s = l.strip()
    if '<w:pPr>' in s:
        in_ppr = True
    elif '</w:pPr>' in s:
        in_ppr = False
    elif in_ppr and ('<w:p>' in s or '<w:p ' in s or '<w:tbl>' in s):
        issues += 1
        print(f'  STILL NESTED at line {i2+1}: {s[:100]}')

print(f'Remaining issues: {issues}')

# Repack
with zipfile.ZipFile(out, 'w', zipfile.ZIP_DEFLATED) as zf:
    for root, dirs, files in os.walk(base):
        for f in files:
            fp = os.path.join(root, f)
            arcname = os.path.relpath(fp, base).replace(os.sep, '/')
            data = open(fp, 'rb').read()
            zf.writestr(arcname, data)

print('Repacked: ' + out)

# Test with python-docx
from docx import Document
try:
    doc = Document(out)
    print('python-docx: OK, paragraphs=' + str(len(doc.paragraphs)))
    for idx, p in enumerate(doc.paragraphs[:3]):
        txt = p.text[:80] if p.text else '(empty)'
        print('  [' + str(idx) + '] ' + txt)
except Exception as e:
    print('python-docx ERROR: ' + str(e))

print('Done!')
