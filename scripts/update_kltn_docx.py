# -*- coding: utf-8 -*-
import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import sys

sys.stdout.reconfigure(encoding='utf-8')

def set_cell_font(cell, text, bold=False, italic=False, font_size=10.5):
    cell.text = text
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for r in p.runs:
            r.font.name = 'Times New Roman'
            r.font.size = Pt(font_size)
            r.font.bold = bold
            r.font.italic = italic

def set_table_borders(table):
    tblPr = table._tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'CCCCCC')
        borders.append(border)
    tblPr.append(borders)

def add_table_after(doc, paragraph, headers, rows):
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)
    
    # Headers
    hdr_cells = table.rows[0].cells
    for i, title in enumerate(headers):
        set_cell_font(hdr_cells[i], title, bold=True, font_size=10.5)
        shading = OxmlElement('w:shd')
        shading.set(qn('w:val'), 'clear')
        shading.set(qn('w:color'), 'auto')
        shading.set(qn('w:fill'), 'EAEAEA')
        hdr_cells[i]._tc.get_or_add_tcPr().append(shading)
        
    # Rows
    for r_idx, row_data in enumerate(rows):
        row_cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row_data):
            set_cell_font(row_cells[c_idx], str(val), bold=False, font_size=10)
            
    paragraph._element.addnext(table._element)
    return table

def find_table_by_header(doc, keyword):
    for table in doc.tables:
        if len(table.rows) > 0:
            hdr_text = " ".join([c.text.strip() for c in table.rows[0].cells])
            if keyword in hdr_text:
                return table
    return None

def main():
    doc_path = r'docs/docs/110122162_TranNhutThien_KLTN.docx'
    print(f"Loading {doc_path}...")
    doc = docx.Document(doc_path)
    
    # -------------------------------------------------------------
    # 1. Update Chapter 1: Section 1.3.2 (Scope)
    # -------------------------------------------------------------
    print("1. Updating Chapter 1 Scope (Section 1.3.2)...")
    for i, p in enumerate(doc.paragraphs):
        if "Phạm vi nghiệp vụ:" in p.text:
            p.text = (
                "Phạm vi nghiệp vụ: Hệ thống tập trung giải quyết các nghiệp vụ cốt lõi xoay quanh 5 vai trò: "
                "Admin (Quản trị hệ thống, duyệt cấp cao nhất, quản lý phân bổ), Cán bộ quỹ (Quản lý thông tin, nhà tài trợ, "
                "duyệt cấp 1 và 2, đề xuất phân bổ và nghiệm thu), Kế toán (Xác nhận dòng tiền vào/ra, giải ngân, quản lý "
                "công nợ và xác nhận thanh toán), Sinh viên (Theo dõi thông tin, nộp đơn trực tuyến và theo dõi lịch trả nợ), "
                "và Ban Kiểm Soát (Giám sát độc lập các hoạt động thu-chi, nghiệm thu, công nợ và lịch sử phê duyệt ở chế độ Read-Only). "
                "Ngoài ra, hệ thống bao gồm đầy đủ 4 phân hệ nghiệp vụ nâng cao: Nghiệm thu hồ sơ dự án, Giám sát công nợ & Lịch trả nợ, "
                "Phân bổ ngân sách nội bộ giữa các quỹ, và Phân quyền kiểm soát độc lập cho Ban Kiểm Soát."
            )
            p.style = 'List Paragraph'
            break

    # -------------------------------------------------------------
    # 2. Update Role 5 in Section 3.1.1 (FR) and Section 3.2 (Actors)
    # -------------------------------------------------------------
    print("2. Adding Role 5 to Section 3.1.1 (FR) and Section 3.2 (Actors)...")
    for i, p in enumerate(doc.paragraphs):
        if "Nhóm chức năng dành cho Quản trị viên hệ thống - Admin" in p.text:
            target_p = doc.paragraphs[i+4]
            
            p1 = target_p.insert_paragraph_before("Nhóm chức năng dành cho Ban Kiểm Soát (Role ID: 5 - Phân loại: BAN_KIEM_SOAT):", style='Normal')
            p2 = target_p.insert_paragraph_before("Tra cứu và Giám sát dòng tiền: Tra cứu danh sách Quỹ, các Khoản tài trợ và Giao dịch Thu/Chi ở chế độ chỉ xem (Read-Only), đảm bảo tính minh bạch tài chính.", style='List Paragraph')
            p3 = target_p.insert_paragraph_before("Giám sát Nghiệm thu & Công nợ: Theo dõi kết quả nghiệm thu hồ sơ (kiểm tra tiến độ và nghiệm thu cuối cùng), tình hình thực hiện hợp đồng vay vốn, điều khoản thu hồi và lịch trả nợ của sinh viên.", style='List Paragraph')
            p4 = target_p.insert_paragraph_before("Xem lịch sử phê duyệt và báo cáo thống kê: Truy vấn toàn bộ lịch sử 3 cấp phê duyệt đơn hỗ trợ, nhật ký hệ thống (Audit Trail) và xem các báo cáo thống kê tổng quan.", style='List Paragraph')
            p5 = target_p.insert_paragraph_before("Ràng buộc quyền hạn: Hệ thống chặn tuyệt đối các thao tác tạo mới, chỉnh sửa, xóa bản ghi, phê duyệt đơn hay thực chi giải ngân đối với vai trò Ban Kiểm Soát.", style='List Paragraph')
            break

    for i, p in enumerate(doc.paragraphs):
        if "Nhà tài trợ(): Người tài trợ cho các quỷ" in p.text or "Quản trị viên (Admin): Người quản lý toàn bộ hệ thống." in p.text:
            if "Ban Kiểm Soát" not in doc.paragraphs[i+1].text:
                doc.paragraphs[i+1].insert_paragraph_before("Ban Kiểm Soát (Supervision Board - Role ID: 5): Thành viên giám sát độc lập toàn bộ hoạt động tài chính, phê duyệt, nghiệm thu và công nợ của hệ thống (quyền Read-Only).", style='Normal')
            break

    # -------------------------------------------------------------
    # 3. Add 3 Use Cases in Section 3.2.2
    # -------------------------------------------------------------
    print("3. Adding 3 Use Cases to Section 3.2.2...")
    for i, p in enumerate(doc.paragraphs):
        if "Mô hình hóa động và Quy trình nghiệp vụ hệ thống (Dynamic Modeling)" in p.text:
            target_p = p
            
            # Use Case 1: Nghiem thu
            target_p.insert_paragraph_before("Đặc tả Use Case: Nghiệm thu hồ sơ dự án / đề xuất hỗ trợ", style='Heading 4')
            target_p.insert_paragraph_before("Tác nhân chính: Cán bộ Quỹ (Role 3), Admin (Role 1), Ban Kiểm Soát (Role 5).", style='Normal')
            target_p.insert_paragraph_before("Tiền điều kiện: Đơn đề nghị hỗ trợ đã được giải ngân (Da gia ngan) và có đánh dấu yêu cầu nghiệm thu (canNghiemThu = 1).", style='List Paragraph')
            target_p.insert_paragraph_before("Luồng chính: 1. Cán bộ Quỹ lập biên bản nghiệm thu (chọn loại: Kiểm tra tiến độ hoặc Nghiệm thu cuối cùng) -> 2. Hệ thống lưu biên bản ở trạng thái 'Chờ đánh giá' -> 3. Admin kiểm tra và đưa ra kết quả phê duyệt ('Đạt', 'Đạt có điều chỉnh', hoặc 'Không đạt') -> 4. Hệ thống cập nhật trạng thái đơn thành 'Da nghiem thu' hoặc 'Nghiem thu khong dat'.", style='List Paragraph')
            target_p.insert_paragraph_before("Hậu điều kiện: Biên bản nghiệm thu được lưu vết vào CSDL và cho phép Ban Kiểm Soát truy vấn giám sát.", style='List Paragraph')
            
            # Use Case 2: Cong no
            target_p.insert_paragraph_before("Đặc tả Use Case: Giám sát và Quản lý Công nợ", style='Heading 4')
            target_p.insert_paragraph_before("Tác nhân chính: Kế toán (Role 2), Admin (Role 1), Cán bộ Quỹ (Role 3), Ban Kiểm Soát (Role 5), Sinh viên (Role 4).", style='Normal')
            target_p.insert_paragraph_before("Tiền điều kiện: Đơn vay vốn hỗ trợ đã được giải ngân và khởi tạo điều khoản thu hồi (dieukhoanthuhoi) cùng lịch trả nợ (lichtrano).", style='List Paragraph')
            target_p.insert_paragraph_before("Luồng chính: 1. Hệ thống tự động theo dõi các kỳ trả nợ đến hạn -> 2. Sinh viên tải lên minh chứng thanh toán -> 3. Kế toán xác nhận minh chứng (chuyển trangthaixacnhan thành 'Da xac nhan' và ghi nhận giao dịch Thu) hoặc từ chối nếu minh chứng sai -> 4. Admin/Cán bộ Quỹ gửi thông báo nhắc nợ đối với các kỳ nợ quá hạn.", style='List Paragraph')
            target_p.insert_paragraph_before("Hậu điều kiện: Số tiền đã thu được cập nhật vào quỹ, giảm dư nợ của sinh viên.", style='List Paragraph')
            
            # Use Case 3: Phan bo
            target_p.insert_paragraph_before("Đặc tả Use Case: Phân bổ Ngân sách nội bộ giữa các Quỹ", style='Heading 4')
            target_p.insert_paragraph_before("Tác nhân chính: Cán bộ Quỹ (Role 3), Admin (Role 1), Ban Kiểm Soát (Role 5).", style='Normal')
            target_p.insert_paragraph_before("Tiền điều kiện: Quỹ nguồn (quỹ cha) có đủ số dư khả dụng.", style='List Paragraph')
            target_p.insert_paragraph_before("Luồng chính: 1. Cán bộ Quỹ tạo đề xuất phân bổ kinh phí từ quỹ cha sang quỹ con -> 2. Hệ thống lưu bản ghi phanbongansach ở trạng thái 'Cho duyet' -> 3. Admin xem xét và thực hiện phê duyệt -> 4. Hệ thống tự động trừ số dư quỹ cha, cộng số dư quỹ con và tạo 2 giao dịch tương ứng (Giao dịch Chi ở quỹ cha, Giao dịch Thu ở quỹ con) -> 5. Admin có thể thực hiện 'Thu hồi' nếu phân bổ bị hủy.", style='List Paragraph')
            target_p.insert_paragraph_before("Hậu điều kiện: Số dư 2 quỹ được điều chỉnh chính xác và lưu vết nhatkyhethong.", style='List Paragraph')
            break

    # -------------------------------------------------------------
    # 4. Add 3 Activity Diagrams in Section 3.3
    # -------------------------------------------------------------
    print("4. Adding 3 Activity Diagrams to Section 3.3...")
    for i, p in enumerate(doc.paragraphs):
        if "Sơ đồ tuần tự (Sequence Diagram) tương tác đa tầng:" in p.text:
            target_p = p
            target_p.insert_paragraph_before("Biểu đồ Activity Diagram 3.3.12: Quy trình Nghiệm thu hồ sơ", style='Heading 4')
            target_p.insert_paragraph_before("Mô tả luồng hoạt động: Cán bộ Quỹ khởi tạo biên bản nghiệm thu -> Hệ thống kiểm tra điều kiện (yeucauhotro đã giải ngân & canNghiemThu=1) -> Lưu trạng thái 'Chờ đánh giá' -> Admin vào thẩm định -> Nếu Đạt: cập nhật đơn thành 'Da nghiem thu'; Nếu Không đạt: cập nhật thành 'Nghiem thu khong dat'.", style='Normal')
            target_p.insert_paragraph_before("[HÌNH ẢNH] Sơ đồ hoạt động Nghiệm thu hồ sơ (AD12_NghiemThu.png)", style='Normal')
            
            target_p.insert_paragraph_before("Biểu đồ Activity Diagram 3.3.13: Quy trình Quản lý và Xác nhận Công nợ", style='Heading 4')
            target_p.insert_paragraph_before("Mô tả luồng hoạt động: Hệ thống tự động khởi tạo Lịch trả nợ khi đơn vay giải ngân -> Sinh viên nộp minh chứng trả nợ -> Kế toán kiểm tra chứng từ -> Nếu khớp: Xác nhận 'Da xac nhan' + Tạo Giao dịch Thu; Nếu không khớp: 'Tu choi' + gửi lý do -> Nếu quá hạn: Admin/Cán bộ gửi thông báo nhắc nợ.", style='Normal')
            target_p.insert_paragraph_before("[HÌNH ẢNH] Sơ đồ hoạt động Công nợ & Lịch trả nợ (AD13_CongNo.png)", style='Normal')
            
            target_p.insert_paragraph_before("Biểu đồ Activity Diagram 3.3.14: Quy trình Phân bổ Ngân sách Quỹ", style='Heading 4')
            target_p.insert_paragraph_before("Mô tả luồng hoạt động: Cán bộ Quỹ tạo yêu cầu phân bổ -> Admin xét duyệt -> Phê duyệt: Trừ tiền Quỹ cha, cộng tiền Quỹ con, cập nhật trạng thái 'Da duyet' -> Từ chối: Lưu lý do từ chối -> Thu hồi: Đảo ngược số dư 2 quỹ về trạng thái ban đầu.", style='Normal')
            target_p.insert_paragraph_before("[HÌNH ẢNH] Sơ đồ hoạt động Phân bổ Ngân sách (AD14_PhanBo.png)", style='Normal')
            break

    # -------------------------------------------------------------
    # 5. Add 3 Sequence Diagrams in Section 3.3.9
    # -------------------------------------------------------------
    print("5. Adding 3 Sequence Diagrams to Section 3.3.9...")
    for i, p in enumerate(doc.paragraphs):
        if "Phân tích các Workflow xử lý đặc thù:" in p.text:
            target_p = p
            target_p.insert_paragraph_before("Sequence Diagram: Quy trình Nghiệm thu hồ sơ", style='Heading 4')
            target_p.insert_paragraph_before("Tương tác giữa các tầng: CanBo -> Frontend UI -> ApplicationController -> NghiemThuModel -> Database (INSERT nghiemthu) -> Admin -> ApplicationController -> Database (UPDATE trangthai yeucauhotro).", style='Normal')
            
            target_p.insert_paragraph_before("Sequence Diagram: Quy trình Quản lý Công nợ", style='Heading 4')
            target_p.insert_paragraph_before("Tương tác giữa các tầng: Student (Upload minh chứng) -> CongNoController -> Database (UPDATE minhchungtrano) -> KeToan -> CongNoController -> Database (UPDATE trangthaixacnhan & INSERT giaodich).", style='Normal')
            
            target_p.insert_paragraph_before("Sequence Diagram: Quy trình Phân bổ Ngân sách", style='Heading 4')
            target_p.insert_paragraph_before("Tương tác giữa các tầng: CanBo -> PhanBoController -> Database (INSERT phanbongansach) -> Admin -> PhanBoController -> Database Transaction (UPDATE sodu quy_cha, UPDATE sodu quy_con, INSERT giaodich).", style='Normal')
            break

    # -------------------------------------------------------------
    # 6. Fill Section 4.2.2 Workflow Engine & State Machine
    # -------------------------------------------------------------
    print("6. Filling Section 4.2.2 (Workflow Engine)...")
    for i, p in enumerate(doc.paragraphs):
        if "Xây dựng công cụ quản lý quy trình phê duyệt (Workflow Engine)" in p.text:
            target_p = doc.paragraphs[i+1]
            
            p1 = target_p.insert_paragraph_before(
                "Phân hệ Workflow Engine chịu trách nhiệm điều phối toàn bộ vòng đời của đơn đề nghị hỗ trợ sinh viên "
                "thông qua Ma trận chuyển trạng thái (State Machine Matrix) và cơ chế phê duyệt 3 cấp nghiêm ngặt.", style='Normal'
            )
            
            p2 = target_p.insert_paragraph_before("1. Ma trận chuyển trạng thái đơn (State Machine Matrix):", style='Normal')
            p3 = target_p.insert_paragraph_before(
                "Luồng chuẩn: [Cho duyet cap 1] -> [Cho duyet cap 2] -> [Cho duyet cap 3] -> [Da gia ngan] -> [Cho nghiem thu] -> [Da nghiem thu].\n"
                "Luồng ngoại lệ: Có thể chuyển sang [Tu choi], [Yeu cau bo sung], [Nghiem thu khong dat], hoặc [Dang thu hoi no] tùy thuộc vào kết quả thẩm định và loại hình hỗ trợ (cấp học bổng / cho vay).", style='Normal'
            )
            
            p4 = target_p.insert_paragraph_before("2. Bảng tổng hợp Enum trangthai trong CSDL (đầy đủ 15+ giá trị):", style='Normal')
            
            headers_enum = ["Trạng thái (Enum)", "Mô tả nghiệp vụ", "Vai trò tác động"]
            rows_enum = [
                ["Cho duyet cap 1", "Đơn mới nộp, chờ Cán bộ Quỹ thẩm định sơ bộ", "Cán bộ Quỹ (Role 3)"],
                ["Cho duyet cap 2", "Đã qua Cấp 1, chờ Admin phê duyệt nội dung", "Admin (Role 1)"],
                ["Cho duyet cap 3", "Đã qua Cấp 2, chờ Kế toán đối soát & thực chi", "Kế toán (Role 2)"],
                ["Da duyet", "Đơn được phê duyệt toàn bộ nhưng chưa giải ngân", "Admin / Kế toán"],
                ["Cho gia ngan", "Chờ quỹ bổ sung số dư để thực chi giải ngân", "Kế toán (Role 2)"],
                ["Da gia ngan", "Đã chuyển tiền giải ngân thành công cho SV", "Kế toán (Role 2)"],
                ["Cho nghiem thu", "Đã giải ngân, chờ Cán bộ lập biên bản nghiệm thu", "Cán bộ Quỹ (Role 3)"],
                ["Da nghiem thu", "Đã nghiệm thu đạt kết quả thành công", "Admin (Role 1)"],
                ["Nghiem thu khong dat", "Nghiệm thu không đạt yêu cầu đề ra", "Admin (Role 1)"],
                ["Dang thu hoi no", "Hợp đồng vay vốn đang trong quá trình trả nợ", "Kế toán / Hệ thống"],
                ["Yeu cau bo sung", "Đơn thiếu minh chứng, yêu cầu SV bổ sung", "Cán bộ / Admin"],
                ["Tu choi", "Đơn bị từ chối phê duyệt ở bất kỳ cấp nào", "Cán bộ / Admin / KT"],
                ["Hoan thanh", "Hoàn tất toàn bộ quy trình đơn & công nợ", "Hệ thống"],
                ["Da huy", "Đơn bị Hủy bởi người nộp hoặc Admin", "Sinh viên / Admin"],
                ["Tam dung", "Tạm dừng xử lý đơn do phát sinh sự cố", "Admin (Role 1)"]
            ]
            add_table_after(doc, target_p, headers_enum, rows_enum)
            
            p5 = target_p.insert_paragraph_before("3. Code snippet minh họa logic chuyển trạng thái Workflow Engine:", style='Normal')
            code_text = (
                "// Middleware / Helper kiểm tra chuyển trạng thái hợp lệ\n"
                "const VALID_TRANSITIONS = {\n"
                "  'Cho duyet cap 1': ['Cho duyet cap 2', 'Tu choi', 'Yeu cau bo sung'],\n"
                "  'Cho duyet cap 2': ['Cho duyet cap 3', 'Tu choi', 'Yeu cau bo sung'],\n"
                "  'Cho duyet cap 3': ['Da gia ngan', 'Cho gia ngan', 'Tu choi'],\n"
                "  'Da gia ngan': ['Cho nghiem thu', 'Dang thu hoi no', 'Hoan thanh'],\n"
                "  'Cho nghiem thu': ['Da nghiem thu', 'Nghiem thu khong dat']\n"
                "};\n\n"
                "function validateStateTransition(currentState, newState) {\n"
                "  const allowed = VALID_TRANSITIONS[currentState] || [];\n"
                "  if (!allowed.includes(newState)) {\n"
                "    throw new Error(`Không thể chuyển trạng thái từ ${currentState} sang ${newState}`);\n"
                "  }\n"
                "  return true;\n"
                "}"
            )
            p_code = target_p.insert_paragraph_before(code_text, style='Normal')
            p_code.runs[0].font.name = 'Courier New'
            p_code.runs[0].font.size = Pt(9.5)
            break

    # -------------------------------------------------------------
    # 7. Update Test Case Table & Section 5.3 (Completed Functions)
    # -------------------------------------------------------------
    print("7. Updating Test Case Table and Section 5.3...")
    tbl25 = find_table_by_header(doc, "Mã TC")
    if tbl25:
        tc_new_rows = [
            ["TC16", "UC06", "Tạo biên bản nghiệm thu", "Đơn Da gia ngan, file PDF hợp lệ", "Tạo nghiệm thu 'Cho danh gia'"],
            ["TC17", "UC06", "Tạo biên bản nghiệm thu", "Đơn chưa giải ngân", "Từ chối, báo lỗi chưa đủ điều kiện"],
            ["TC18", "UC06", "Duyệt nghiệm thu", "Admin duyệt kết quả 'Dat'", "Cập nhật đơn 'Da nghiem thu'"],
            ["TC19", "UC06", "Duyệt nghiệm thu", "Admin duyệt kết quả 'Khong dat'", "Cập nhật đơn 'Nghiem thu khong dat'"],
            ["TC20", "UC06", "Sửa biên bản nghiệm thu", "Cán bộ sửa khi chưa duyệt", "Cập nhật thông tin biên bản"],
            ["TC21", "UC07", "Tạo lịch trả nợ", "Hợp đồng vay giải ngân thành công", "Tự động sinh các kỳ trả nợ"],
            ["TC22", "UC07", "Xác nhận trả nợ", "Kế toán duyệt minh chứng trả nợ", "Cập nhật 'Da xac nhan' & Giao dịch Thu"],
            ["TC23", "UC07", "Nhắc nợ quá hạn", "Hệ thống quét khoản nợ quá hạn", "Gửi thông báo nhắc nợ cho SV"],
            ["TC24", "UC08", "Đề xuất phân bổ", "Cán bộ đề xuất trích từ quỹ cha", "Tạo bản ghi phanbongansach Cho duyet"],
            ["TC25", "UC08", "Duyệt/Thu hồi phân bổ", "Admin duyệt hoặc thu hồi phân bổ", "Điều chỉnh số dư 2 quỹ & ghi log"],
            ["TC26", "UC09", "Kiểm soát Role 5", "Role 5 truy cập API & Dashboard", "Xem OK (Read-Only), Chặn 403 khi Mutation"]
        ]
        for r_data in tc_new_rows:
            row = tbl25.add_row()
            for idx, val in enumerate(r_data):
                set_cell_font(row.cells[idx], val, font_size=10)

    for i, p in enumerate(doc.paragraphs):
        if "Chức năng đã hoàn thành" in p.text:
            target_p = doc.paragraphs[i+1]
            target_p.insert_paragraph_before("5. Phân hệ Nghiệm thu hồ sơ (Kiểm tra tiến độ & Nghiệm thu cuối cùng).", style='List Paragraph')
            target_p.insert_paragraph_before("6. Phân hệ Quản lý Công nợ, Hợp đồng vay vốn & Lịch trả nợ.", style='List Paragraph')
            target_p.insert_paragraph_before("7. Phân hệ Phân bổ Ngân sách nội bộ giữa các Quỹ.", style='List Paragraph')
            target_p.insert_paragraph_before("8. Phân quyền và Màn hình giám sát cho Ban Kiểm Soát (Role 5 - Read-Only).", style='List Paragraph')
            break

    # -------------------------------------------------------------
    # 8. Update Chapter 6: Section 6.1 (Conclusion) & 6.3 (Future Work)
    # -------------------------------------------------------------
    print("8. Updating Chapter 6 (Conclusion & Future Work)...")
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == "Kết luận":
            target_p = doc.paragraphs[i+1]
            target_p.insert_paragraph_before(
                "Đề tài 'Xây dựng Hệ thống Quản lý Quỹ Phát triển Đại học Trà Vinh (TVU Fund Management)' đã hoàn thành xuất sắc "
                "các mục tiêu nghiên cứu và triển khai thực tế. Hệ thống được hoàn thiện toàn diện với 5 phân hệ vai trò (Admin, Kế toán, "
                "Cán bộ Quỹ, Sinh viên, Ban Kiểm Soát), cơ sở dữ liệu chuẩn hóa gồm 21+ bảng thực thể, 8+ nhóm Use Cases cốt lõi cùng các "
                "workflow xử lý tài chính phức tạp bao gồm phê duyệt 3 cấp, phân bổ ngân sách, nghiệm thu hồ sơ và quản lý công nợ.", style='Normal'
            )
            break

    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == "Hướng phát triển":
            target_p = doc.paragraphs[i+1]
            target_p.insert_paragraph_before(
                "Lưu ý: Các tính năng nâng cao bao gồm Nghiệm thu hồ sơ, Quản lý công nợ & Lịch trả nợ, Phân bổ ngân sách giữa các quỹ "
                "và Phân quyền Ban Kiểm Soát (Role 5) ban đầu dự kiến thuộc hướng phát triển nhưng ĐÃ ĐƯỢC TRIỂN KHAI HOÀN THIỆN trong phiên bản này.", style='Normal'
            )
            break

    # -------------------------------------------------------------
    # 9. Update Phụ Lục A (Appendix A Data Dictionary)
    # -------------------------------------------------------------
    print("9. Updating Phụ Lục A (Data Dictionary)...")
    for i, p in enumerate(doc.paragraphs):
        if "Phụ lục A: Data Dictionary đầy đủ" in p.text:
            target_p = doc.paragraphs[i+1]
            
            # Table 1: nghiemthu
            target_p.insert_paragraph_before("A.nghiemthu — Bảng Nghiệm thu hồ sơ dự án", style='Heading 3')
            h1 = ["Cột", "Kiểu dữ liệu / Ràng buộc", "Ghi chú"]
            r1 = [
                ["nghiemthu_id", "int(11) AUTO_INCREMENT PRIMARY KEY", "Mã nghiệm thu"],
                ["yeucauhotro_id", "int(11) NOT NULL FK -> yeucauhotro", "Mã đơn hỗ trợ cần nghiệm thu"],
                ["lanthu", "int(11) DEFAULT 1", "Lần nghiệm thu (1, 2...)"],
                ["loaikiemtra", "enum('Kiem tra tien do', 'Nghiem thu cuoi cung')", "Loại kiểm tra"],
                ["ketqua", "enum('Cho danh gia', 'Dat', 'Dat co dieu chinh', 'Khong dat')", "Kết quả thẩm định"],
                ["soquyetdinh", "varchar(100) NULL", "Số quyết định thành lập hội đồng"],
                ["filebienban", "varchar(255) NULL", "Đường dẫn file biên bản PDF"],
                ["nguoinghiemthu_id", "int(11) NULL FK -> nguoidung", "Cán bộ / Admin thực hiện"],
                ["nhanxet", "text NULL", "Nhận xét đánh giá của hội đồng"],
                ["ngaynghiemthu", "timestamp DEFAULT CURRENT_TIMESTAMP", "Ngày thực hiện nghiệm thu"]
            ]
            add_table_after(doc, target_p, h1, r1)
            
            # Table 2: lichtrano
            target_p.insert_paragraph_before("A.lichtrano — Bảng Lịch trả nợ khoản vay hỗ trợ", style='Heading 3')
            r2 = [
                ["lichtrano_id", "int(11) AUTO_INCREMENT PRIMARY KEY", "Mã lịch trả nợ"],
                ["dieukhoanthuhoi_id", "int(11) NOT NULL FK -> dieukhoanthuhoi", "Mã điều khoản thu hồi"],
                ["ky", "int(11) NOT NULL", "Kỳ trả nợ (1, 2, 3...)"],
                ["ngaydenhan", "date NOT NULL", "Ngày đến hạn thanh toán"],
                ["sotienphaitra", "decimal(15,2) NOT NULL", "Số tiền phải trả trong kỳ"],
                ["trangthai", "enum('Chua thanh toan', 'Da thanh toan', 'Qua han')", "Trạng thái kỳ nợ"],
                ["trangthaixacnhan", "enum('Chua xac nhan', 'Da xac nhan', 'Tu choi')", "Kế toán xác nhận minh chứng"],
                ["ngayxacnhan", "timestamp NULL", "Ngày Kế toán duyệt minh chứng"],
                ["nguoiduyet_id", "int(11) NULL FK -> nguoidung", "Kế toán duyệt"],
                ["minhchungtrano", "varchar(255) NULL", "File minh chứng chuyển khoản của SV"],
                ["ghichuxacnhan", "text NULL", "Ghi chú của Kế toán khi xác nhận/từ chối"]
            ]
            add_table_after(doc, target_p, h1, r2)
            
            # Table 3: dieukhoanthuhoi
            target_p.insert_paragraph_before("A.dieukhoanthuhoi — Bảng Điều khoản thu hồi vốn vay", style='Heading 3')
            r3 = [
                ["dieukhoanthuhoi_id", "int(11) AUTO_INCREMENT PRIMARY KEY", "Mã điều khoản thu hồi"],
                ["yeucauhotro_id", "int(11) NOT NULL FK -> yeucauhotro", "Mã đơn vay hỗ trợ"],
                ["tongsotien", "decimal(15,2) NOT NULL", "Tổng số tiền cần hoàn trả"],
                ["phantram", "decimal(5,2) DEFAULT 100.00", "% kinh phí phải hoàn trả"],
                ["sotienthucte", "decimal(15,2) NOT NULL", "Số tiền thực tế tính theo %"],
                ["ngaybatdauthuhoi", "date NULL", "Ngày bắt đầu tính lịch thu hồi"],
                ["sotiendadathu", "decimal(15,2) DEFAULT 0.00", "Tổng số tiền đã thu lại được"],
                ["trangthai", "enum('Dang thu hoi', 'Hoan thanh', 'Qua han')", "Trạng thái điều khoản thu hồi"]
            ]
            add_table_after(doc, target_p, h1, r3)

            # Table 4: dotgiaingan
            target_p.insert_paragraph_before("A.dotgiaingan — Bảng Đợt giải ngân theo tiến độ", style='Heading 3')
            r4 = [
                ["dotgiaingan_id", "int(11) AUTO_INCREMENT PRIMARY KEY", "Mã đợt giải ngân"],
                ["yeucauhotro_id", "int(11) NOT NULL FK -> yeucauhotro", "Mã đơn hỗ trợ"],
                ["dot", "int(11) NOT NULL", "Số đợt (Đợt 1, Đợt 2...)"],
                ["sotien", "decimal(15,2) NOT NULL", "Số tiền giải ngân trong đợt"],
                ["ngaygiaingan", "timestamp DEFAULT CURRENT_TIMESTAMP", "Ngày thực hiện giải ngân"],
                ["ngaybatdau", "date NULL", "Ngày bắt đầu giai đoạn giải ngân"],
                ["ngayketthuc", "date NULL", "Ngày kết thúc giai đoạn giải ngân"]
            ]
            add_table_after(doc, target_p, h1, r4)
            break

    # -------------------------------------------------------------
    # 10. Update Phụ Lục D (Role Guide & API Table)
    # -------------------------------------------------------------
    print("10. Updating Phụ Lục D (Role Guide & API Table)...")
    tbl45 = find_table_by_header(doc, "Thao tác chính")
    if tbl45:
        set_cell_font(tbl45.rows[3].cells[1], "Quản lý quỹ, duyệt cấp 1, đề xuất phân bổ ngân sách, lập biên bản nghiệm thu hồ sơ, quản lý nhà tài trợ/người dùng/nội dung")
        r_bks = tbl45.add_row()
        set_cell_font(r_bks.cells[0], "Ban Kiểm Soát (Role 5)", bold=True)
        set_cell_font(r_bks.cells[1], "Giám sát độc lập dữ liệu Quỹ, Khoản tài trợ, Giao dịch Thu/Chi, Lịch sử phê duyệt 3 cấp, Biên bản nghiệm thu & Công nợ (quyền Read-Only)")

    tbl46 = find_table_by_header(doc, "Endpoint")
    if tbl46:
        api_new_rows = [
            ["GET", "/api/bks/dashboard", "Role 5", "Dashboard tổng quan giám sát tài chính dành cho Ban Kiểm Soát"],
            ["GET", "/api/bks/nghiem-thu", "Role 5", "Xem danh sách và chi tiết biên bản nghiệm thu (Read-Only)"],
            ["GET", "/api/bks/cong-no", "Role 5", "Xem danh sách hợp đồng công nợ & lịch trả nợ (Read-Only)"],
            ["GET", "/api/nghiem-thu", "Role 1,3,5", "Danh sách biên bản nghiệm thu hồ sơ"],
            ["POST", "/api/nghiem-thu", "Role 1,3", "Tạo mới biên bản nghiệm thu hồ sơ"],
            ["PUT", "/api/nghiem-thu/:id/duyet", "Role 1", "Admin phê duyệt kết quả nghiệm thu (Đạt/Không đạt)"],
            ["GET", "/api/cong-no/lich-tra-no", "Role 1,2,3,4,5", "Xem lịch trả nợ (SV xem của mình, KT/BKS xem toàn hệ thống)"],
            ["PUT", "/api/cong-no/xac-nhan", "Role 2", "Kế toán xác nhận minh chứng trả nợ & ghi nhận giao dịch Thu"],
            ["POST", "/api/phan-bo", "Role 1,3", "Cán bộ đề xuất phân bổ ngân sách từ Quỹ cha sang Quỹ con"],
            ["PUT", "/api/phan-bo/:id/duyet", "Role 1", "Admin phê duyệt đề xuất phân bổ & tự động chuyển số dư"]
        ]
        for r_data in api_new_rows:
            row = tbl46.add_row()
            for idx, val in enumerate(r_data):
                set_cell_font(row.cells[idx], val, font_size=10)

    # -------------------------------------------------------------
    # Save Output
    # -------------------------------------------------------------
    print("Saving updated Word document...")
    doc.save(doc_path)
    print("DONE! Successfully updated 110122162_TranNhutThien_KLTN.docx")

if __name__ == '__main__':
    main()
